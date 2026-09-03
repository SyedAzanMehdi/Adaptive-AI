"""Convert markdown docs to Word (.docx) with clean professional styling.

Usage: python scripts/md_to_docx.py <input.md> <output.docx> [<input.md> <output.docx> ...]
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

DARK = RGBColor(0x11, 0x11, 0x11)
GRAY = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0x00, 0x00, 0x00)

INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def add_runs(paragraph, text, base_size=11, bold_all=False, color=DARK):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
            run.bold = bold_all
        run.font.size = Pt(base_size)
        run.font.color.rgb = color
        if run.font.name is None:
            run.font.name = "Calibri"


def shade_cell(cell, fill="E7E6E6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): fill})
    tc_pr.append(shd)


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            cell = table.cell(i, j)
            text = row[j] if j < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, text, base_size=10, bold_all=(i == 0))
            if i == 0:
                shade_cell(cell)
    doc.add_paragraph()


def convert(md_path, out_path):
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    n = len(lines)
    first_h1 = True
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            i += 1
            code = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1  # closing fence
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            p.paragraph_format.left_indent = Inches(0.3)
            continue

        if stripped.startswith("|"):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells if c):
                    rows.append(cells)
                i += 1
            add_table(doc, rows)
            continue

        if stripped in ("---", "***", "___"):
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            if level == 1:
                h = doc.add_heading(level=0 if first_h1 else 1)
                first_h1 = False
            else:
                h = doc.add_heading(level=level - 1)
            add_runs(h, text, base_size={1: 20, 2: 15, 3: 12.5, 4: 11.5}[level], bold_all=True, color=ACCENT)
            i += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            text = re.sub(r"^[-*]\s+", "", stripped)
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, text)
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(2))
            i += 1
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, stripped.lstrip("> "), color=GRAY)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        if stripped == "":
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(out_path)
    print(f"OK {md_path} -> {out_path}")


def main():
    args = sys.argv[1:]
    if len(args) % 2 != 0 or not args:
        print(__doc__)
        sys.exit(1)
    for k in range(0, len(args), 2):
        convert(args[k], args[k + 1])


if __name__ == "__main__":
    main()
